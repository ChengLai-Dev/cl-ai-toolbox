# -*- coding: utf-8 -*-

import json
import os
import re
from typing import Any, Dict, List, Optional
from copy import deepcopy

from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent


def _runs_text(p) -> str:
    """获取段落跨 run 的完整文本"""
    return "".join(r.text for r in p.runs)


def _clear_and_set_text(p, text):
    """清除段落所有 run，写入新文本，保留第一个 run 的格式"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rpr = None
    if p.runs:
        rpr_elem = p.runs[0]._element.find(qn('w:rPr'))
        if rpr_elem is not None:
            rpr = deepcopy(rpr_elem)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    new_run = OxmlElement('w:r')
    if rpr is not None:
        new_run.append(rpr)
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_run.append(new_t)
    p._element.append(new_run)


def _read_docx(file_path: str) -> Dict:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            tables.append(rows)
        return {
            "success": True,
            "paragraphs": paragraphs,
            "tables": tables,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def _read_doc(file_path: str) -> Dict:
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(file_path))
            content = doc.Content.Text
            doc.Close()
            return {"success": True, "text": content}
        finally:
            word.Quit()
    except Exception as e:
        return {"success": False, "error": f"无法读取 .doc 文件（需要安装 pywin32 且系统需安装 Word）: {e}"}


_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.tif', '.webp'}
_IMAGE_MARKER_RE = re.compile(r'\{image:([^}]+)\}')


def _parse_image_markers(text: str) -> list[dict]:
    """解析段落中所有 {image:path} 标记，返回混合片段列表。
    每段格式：{'type': 'text', 'text': str} 或 {'type': 'image', 'path': str}
    若图片文件不存在，标记原样保留为文本。
    """
    segments = []
    last_end = 0
    for m in _IMAGE_MARKER_RE.finditer(text):
        if m.start() > last_end:
            segments.append({'type': 'text', 'text': text[last_end:m.start()]})
        path = m.group(1).strip()
        ext = os.path.splitext(path)[1].lower()
        if ext in _IMAGE_EXTENSIONS and os.path.exists(path):
            segments.append({'type': 'image', 'path': path})
        else:
            segments.append({'type': 'text', 'text': m.group(0)})
        last_end = m.end()
    if last_end < len(text):
        segments.append({'type': 'text', 'text': text[last_end:]})
    return segments


def _build_paragraph(para_text: str, doc):
    """根据段落文本构建段落，支持文本与 {image:path} 图片混合。"""
    from docx.shared import Inches
    segments = _parse_image_markers(para_text)
    has_image = any(s['type'] == 'image' for s in segments)
    if not has_image:
        doc.add_paragraph(para_text)
        return
    p = doc.add_paragraph()
    for seg in segments:
        if seg['type'] == 'text':
            p.add_run(seg['text'])
        else:
            p.add_run().add_picture(seg['path'], width=Inches(5.5))


def _create_docx(file_path: str, paragraphs: List[str]) -> Dict:
    try:
        from docx import Document
        doc = Document()
        for para in paragraphs:
            _build_paragraph(para, doc)
        os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
        doc.save(file_path)
        return {"success": True, "message": f"文档已保存到: {file_path}", "path": file_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _insert_image(file_path: str, image_path: str, paragraph_index: int = -1,
                  width_inches: float = 5.5) -> Dict:
    try:
        ext = os.path.splitext(image_path)[1].lower()
        if ext not in _IMAGE_EXTENSIONS or not os.path.exists(image_path):
            return {"success": False, "error": f"图片文件不存在或格式不支持: {image_path}"}
        from docx import Document
        from docx.shared import Inches
        doc = Document(file_path)
        if paragraph_index >= 0 and paragraph_index < len(doc.paragraphs):
            ref_p = doc.paragraphs[paragraph_index]
            ref_p.add_run().add_picture(image_path, width=Inches(width_inches))
            inserted_at = paragraph_index
        else:
            doc.add_picture(image_path, width=Inches(width_inches))
            inserted_at = len(doc.paragraphs) - 1
        doc.save(file_path)
        return {"success": True, "message": f"图片已插入到位置 {inserted_at}",
                "path": file_path, "inserted_at": inserted_at}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _set_paragraph_text(file_path: str, index: int, text: str) -> Dict:
    try:
        from docx import Document
        doc = Document(file_path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "error": f"段落索引 {index} 超出范围（0-{len(doc.paragraphs)-1}）"}
        _clear_and_set_text(doc.paragraphs[index], text)
        doc.save(file_path)
        return {"success": True, "message": f"段落[{index}] 文本已设置", "path": file_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _insert_text_docx(file_path: str, text: str, paragraph_index: int = -1) -> Dict:
    try:
        from docx import Document
        doc = Document(file_path)
        if paragraph_index >= 0 and paragraph_index < len(doc.paragraphs):
            doc.paragraphs[paragraph_index].add_run(text)
        else:
            doc.add_paragraph(text)
        doc.save(file_path)
        return {"success": True, "message": "文本已插入", "path": file_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _replace_text_docx(file_path: str, old_text: str, new_text: str) -> Dict:
    try:
        from docx import Document
        doc = Document(file_path)
        count = 0
        for para in doc.paragraphs:
            full = _runs_text(para)
            if old_text in full:
                _clear_and_set_text(para, full.replace(old_text, new_text))
                count += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full = _runs_text(para)
                        if old_text in full:
                            _clear_and_set_text(para, full.replace(old_text, new_text))
                            count += 1
        doc.save(file_path)
        return {"success": True, "message": f"替换完成，共替换 {count} 处", "path": file_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _remove_paragraphs(file_path: str, indices: List[int] = None, text_contains: str = None) -> Dict:
    try:
        from docx import Document
        doc = Document(file_path)
        removed = []
        paras_to_remove = []

        if indices:
            for idx in sorted(indices, reverse=True):
                if 0 <= idx < len(doc.paragraphs):
                    paras_to_remove.append(doc.paragraphs[idx])
                    removed.append({"index": idx, "text": doc.paragraphs[idx].text})

        if text_contains:
            for idx in range(len(doc.paragraphs) - 1, -1, -1):
                if text_contains in doc.paragraphs[idx].text:
                    if doc.paragraphs[idx] not in paras_to_remove:
                        paras_to_remove.append(doc.paragraphs[idx])
                        removed.append({"index": idx, "text": doc.paragraphs[idx].text, "matched_by": "text_contains"})

        for p in paras_to_remove:
            p._element.getparent().remove(p._element)

        doc.save(file_path)
        return {"success": True, "message": f"已删除 {len(removed)} 个段落", "removed": removed, "path": file_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


def _insert_paragraph(file_path: str, text: str, index: int = -1) -> Dict:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
        doc = Document(file_path)
        paras = doc.paragraphs
        if index < 0 or index >= len(paras):
            doc.add_paragraph(text)
            inserted_at = len(doc.paragraphs) - 1
        else:
            ref_p = paras[index]
            new_p = OxmlElement('w:p')
            ref_p._element.addnext(new_p)
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, ref_p._element.getparent())
            _clear_and_set_text(new_para, text)
            inserted_at = index + 1
        doc.save(file_path)
        return {"success": True, "message": f"段落已插入到位置 {inserted_at}", "path": file_path, "inserted_at": inserted_at}
    except Exception as e:
        return {"success": False, "error": str(e)}


server = Server("doc-mcp")


@server.list_tools()
async def list_tools() -> List[Tool]:
    return [
        Tool(
            name="read_doc",
            description="读取 Word 文档内容（支持 .docx 和 .doc 格式）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Word 文档的绝对路径"}
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="create_docx",
            description="创建一个新的 .docx 文件。若段落以 {image:路径} 格式标记则嵌入为图片，否则视为文本段落",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "要创建的 .docx 文件路径"},
                    "paragraphs": {"type": "array", "items": {"type": "string"}, "description": "段落列表。文本直接写字符串；若需要插入图片，用 {image:图片绝对路径} 格式标记"}
                },
                "required": ["file_path", "paragraphs"]
            }
        ),
        Tool(
            name="set_paragraph_text",
            description="设置指定索引段落的完整文本（跨 run 合并后写入）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "index": {"type": "integer", "description": "段落索引（从0开始）"},
                    "text": {"type": "string", "description": "新的段落文本"}
                },
                "required": ["file_path", "index", "text"]
            }
        ),
        Tool(
            name="insert_text_docx",
            description="向已有 .docx 文件中的指定段落追加文本",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "text": {"type": "string", "description": "要插入的文本"},
                    "paragraph_index": {"type": "integer", "description": "段落索引（从0开始），-1 表示追加到末尾", "default": -1}
                },
                "required": ["file_path", "text"]
            }
        ),
        Tool(
            name="insert_paragraph",
            description="在指定位置插入一个新段落",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "text": {"type": "string", "description": "新段落的文本"},
                    "index": {"type": "integer", "description": "在此索引后插入，-1 表示追加到末尾", "default": -1}
                },
                "required": ["file_path", "text"]
            }
        ),
        Tool(
            name="replace_text_docx",
            description="替换 .docx 文件中的文本内容（跨 run 合并识别，保留第一个 run 的格式）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "old_text": {"type": "string", "description": "要查找的旧文本"},
                    "new_text": {"type": "string", "description": "替换后的新文本"}
                },
                "required": ["file_path", "old_text", "new_text"]
            }
        ),
        Tool(
            name="remove_paragraphs",
            description="删除段落（按索引列表或按文本内容匹配）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "indices": {"type": "array", "items": {"type": "integer"}, "description": "要删除的段落索引列表（可选）"},
                    "text_contains": {"type": "string", "description": "删除包含此文本的所有段落（可选）"}
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="insert_image",
            description="向已有 .docx 文件的指定位置插入一张图片",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": ".docx 文件路径"},
                    "image_path": {"type": "string", "description": "图片文件的绝对路径（支持 png/jpg/bmp/gif/tiff/webp）"},
                    "paragraph_index": {"type": "integer", "description": "在此段落追加图片，-1 表示追加到末尾", "default": -1},
                    "width_inches": {"type": "number", "description": "图片宽度（英寸），默认 5.5", "default": 5.5}
                },
                "required": ["file_path", "image_path"]
            }
        )
    ]


@server.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
    result = {}

    if name == "read_doc":
        file_path = arguments.get("file_path", "")
        if not file_path or not os.path.exists(file_path):
            result = {"success": False, "error": f"文件不存在: {file_path}"}
        else:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".docx":
                result = _read_docx(file_path)
            elif ext == ".doc":
                result = _read_doc(file_path)
            else:
                result = {"success": False, "error": f"不支持的文件格式: {ext}，仅支持 .doc 和 .docx"}

    elif name == "create_docx":
        file_path = arguments.get("file_path", "")
        paragraphs = arguments.get("paragraphs", [])
        result = _create_docx(file_path, paragraphs) if file_path else {"success": False, "error": "文件路径不能为空"}

    elif name == "set_paragraph_text":
        file_path = arguments.get("file_path", "")
        index = arguments.get("index", -1)
        text = arguments.get("text", "")
        if not file_path or index < 0:
            result = {"success": False, "error": "文件路径和索引不能为空"}
        elif not os.path.exists(file_path):
            result = {"success": False, "error": f"文件不存在: {file_path}"}
        else:
            result = _set_paragraph_text(file_path, index, text)

    elif name == "insert_text_docx":
        file_path = arguments.get("file_path", "")
        text = arguments.get("text", "")
        paragraph_index = arguments.get("paragraph_index", -1)
        if not file_path or not text:
            result = {"success": False, "error": "文件路径和文本不能为空"}
        elif not os.path.exists(file_path):
            result = {"success": False, "error": f"文件不存在: {file_path}"}
        else:
            result = _insert_text_docx(file_path, text, paragraph_index)

    elif name == "insert_paragraph":
        file_path = arguments.get("file_path", "")
        text = arguments.get("text", "")
        index = arguments.get("index", -1)
        if not file_path or not text:
            result = {"success": False, "error": "文件路径和文本不能为空"}
        elif not os.path.exists(file_path):
            result = {"success": False, "error": f"文件不存在: {file_path}"}
        else:
            result = _insert_paragraph(file_path, text, index)

    elif name == "replace_text_docx":
        file_path = arguments.get("file_path", "")
        old_text = arguments.get("old_text", "")
        new_text = arguments.get("new_text", "")
        if not file_path or not old_text:
            result = {"success": False, "error": "文件路径和旧文本不能为空"}
        elif not os.path.exists(file_path):
            result = {"success": False, "error": f"文件不存在: {file_path}"}
        else:
            result = _replace_text_docx(file_path, old_text, new_text)

    elif name == "remove_paragraphs":
        file_path = arguments.get("file_path", "")
        indices = arguments.get("indices")
        text_contains = arguments.get("text_contains")
        if not file_path:
            result = {"success": False, "error": "文件路径不能为空"}
        elif not os.path.exists(file_path):
            result = {"success": False, "error": f"文件不存在: {file_path}"}
        elif not indices and not text_contains:
            result = {"success": False, "error": "请提供 indices 或 text_contains 参数"}
        else:
            result = _remove_paragraphs(file_path, indices, text_contains)

    elif name == "insert_image":
        file_path = arguments.get("file_path", "")
        image_path = arguments.get("image_path", "")
        paragraph_index = arguments.get("paragraph_index", -1)
        width_inches = arguments.get("width_inches", 5.5)
        if not file_path or not image_path:
            result = {"success": False, "error": "file_path 和 image_path 不能为空"}
        elif not os.path.exists(file_path):
            result = {"success": False, "error": f"目标文档不存在: {file_path}"}
        else:
            result = _insert_image(file_path, image_path, paragraph_index, width_inches)

    else:
        result = {"success": False, "error": f"未知的工具: {name}"}

    return [TextContent(type="text", text=json.dumps(result, ensure_ascii=False, indent=2))]


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream, server.create_initialization_options())


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
